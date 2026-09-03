#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成数模解题标准目录骨架 + 锁版本 requirements + export_results.py 模板 + 论文模板（LaTeX/DOCX）。

用法:
    python init_project.py <题名缩写> [--dir <路径>] [--template {cumcm,mcm,both}]

示例:
    python init_project.py cumcm2026A
    python init_project.py mcm2027C --dir ./work --template mcm
    python init_project.py test2025 --template both

生成结构:
    <题名>/
      data/           原始数据
      src/            q1.py q2.py ... + export_results.py
      out/            results.json(空模板) + 图表(png/svg,>=300dpi)
      report/         main.tex / main.docx + figures/
      requirements.txt  锁版本(==)
      README.md
"""
import argparse
import json
import os
import sys

EXPORT_TEMPLATE = '''#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""把全部关键数字落成 out/results.json（数值单一来源，论文只引用它）。

用法: 各 qN.py 算完把结果 import 进来，或在本文件里直接 import 各模块后汇总。
原则: 论文/摘要/图注里的每个数字都必须能在这个文件里找到出处。
"""
import json
import os

# 把所有要写进论文的数字放进这个字典（含单位/置信，不要手抄，用代码算出来填）
RESULTS = {
    # "q1_optimal_value": 123.45,
    # "q1_error_pct": 2.3,
}


def dump(path="out/results.json"):
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(RESULTS, f, ensure_ascii=False, indent=2)
    print(f"[export_results] 已写出 {len(RESULTS)} 个关键数字 -> {path}")


if __name__ == "__main__":
    # TODO: 在此 import 各 qN 模块，把它们的计算结果填进 RESULTS
    dump()
'''

README_TEMPLATE = """# {name}

生成时间: {ts}

## 目录约定
- `data/` 原始数据（不要直接改，清洗产物放 `out/clean.csv`）
- `src/` 每个子问题的求解脚本 `q1.py q2.py ...`，统一由 `export_results.py` 导出数字
- `out/` 结果表格、图（png/svg，>=300dpi）、`results.json`
- `report/` 论文 `main.tex`/`main.docx`、图表 `figures/`

## 铁律
1. 论文每个数字必须来自 `out/results.json`（数值单一来源）。
2. 固定随机种子，交付前 canonical 脚本复跑逐位核对。
3. 改一个数字必 grep 全目录旧值。
"""

REQUIREMENTS = """# 占位版本（并非常年最优）：交付前在本项目 venv 里 `pip freeze` 后精修为真实锁定版，见 references/reproducibility.md 铁律二。
# 锁版本，精确到 == ，干净环境 pip install -r 可复跑
numpy==1.26.4
scipy==1.13.1
pandas==2.2.2
scikit-learn==1.5.1
matplotlib==3.9.2
statsmodels==0.14.2
"""

# LaTeX 模板（国赛 CUMCM，中文，GB/T 7714）
LATEX_TEMPLATE = r"""
\documentclass[12pt,a4paper]{ctexart}
\usepackage{geometry}
\usepackage{amsmath,amssymb,booktabs,graphicx,float,caption,hyperref}
\usepackage{listings,xcolor,indentfirst,array,multirow,makecell}
\usepackage{algorithm,algorithmicx,algpseudocode}
\usepackage{siunitx}
\usepackage{titlesec}
\usepackage{setspace}

% 页面设置
\geometry{left=2.54cm,right=2.54cm,top=2.54cm,bottom=2.54cm}
\setstretch{1.25}
\titleformat{\section}{\centering\heiti\zihao{-3}}{\thesection}{1em}{}
\titleformat{\subsection}{\heiti\zihao{4}}{\thesubsection}{1em}{}
\renewcommand{\theequation}{\arabic{section}.\arabic{equation}}

% 代码块
\lstset{
    basicstyle=\ttfamily\small,
    keywordstyle=\color{blue},
    commentstyle=\color{green!60!black},
    stringstyle=\color{red},
    showstringspaces=false,
    breaklines=true,
    frame=single,
    numbers=left,
    numberstyle=\tiny,
}

% 超链接
\hypersetup{
    colorlinks=true,
    linkcolor=black,
    citecolor=blue,
    urlcolor=blue
}

% 三线表命令
\newcommand{\tabincell}[2]{\begin{tabular}{@{}#1@{}}#2\end{tabular}}

\begin{document}

\title{基于\textbf{<模型名>}的\textbf{<对象>}\textbf{<任务>}}
\author{作者名}
\date{\today}
\maketitle

\begin{abstract}
\noindent \textbf{摘要}

本文针对<问题一句话>。在第一问中，构建了<模型名>，通过<关键动作>得出<量化结论>；第二问中……（依次覆盖每问）。结果表明<带具体数值的结论>，<灵敏度表现>，<可推广方向>。

\textbf{关键词：} 关键词1；关键词2；关键词3
\end{abstract}

\section{问题重述}
用自己的话重述题目背景与各问要求，不抄题面原文。

\section{问题分析}
逐问分析：问题类型、数据字段匹配、初步建模思路、评价指标。

\section{模型假设}
\begin{enumerate}
    \item 假设1：<具体可辩护的判断>。理由：<理由/文献>。失效影响：<若假设不成立会怎样>。
    \item 假设2：……
\end{enumerate}

\section{符号说明}
\begin{table}[H]
\centering
\caption{主要符号表}
\label{tab:symbols}
\begin{tabular}{cll}
\toprule
符号 & 含义 & 单位/备注 \\
\midrule
$N$ & 样本数 & -- \\
$X$ & 特征矩阵 & -- \\
$y$ & 目标变量 & -- \\
\bottomrule
\end{tabular}
\end{table}

\section{模型建立与求解}
\subsection{问题一：<问题名>}
\subsubsection{模型构建}
<模型推导过程，给出核心公式，模型要有名字如 Model I / FMS / RFMS 等>

\subsubsection{求解算法}
<算法描述 + 伪代码或关键代码片段>

\subsubsection{结果}
\begin{table}[H]
\centering
\caption{问题一关键结果}
\label{tab:q1_result}
\begin{tabular}{lcc}
\toprule
指标 & 数值 & 单位 \\
\midrule
最优值 & \texttt{\textbackslash{}py\{q1_optimal_value\}} & -- \\
相对误差 & \texttt{\textbackslash{}py\{q1_error_pct\}} & \% \\
\bottomrule
\end{tabular}
\end{table}

% 结果图示例
%\begin{figure}[H]
%\centering
%\includegraphics[width=0.9\linewidth]{figures/q1_result.png}
%\caption{问题一结果图}
%\label{fig:q1_result}
%\end{figure}

\subsection{问题二：<问题名>}
……

\subsection{问题三：<问题名>}
……

\section{模型检验与灵敏度分析}
\subsection{量纲/量级/边界校验}
已通过 \texttt{sanity\_check.py} 自动校验。

\subsection{灵敏度分析}
对关键参数进行 $\pm 10\%$ / $\pm 20\%$ 摄动，观察结果变化。
\begin{table}[H]
\centering
\caption{灵敏度分析结果}
\label{tab:sensitivity}
\begin{tabular}{lccc}
\toprule
参数 & 摄动幅度 & 结果变化率 & 结论 \\
\midrule
权重 $w$ & $\pm 10\%$ & $<2\%$ & 稳健 \\
\bottomrule
\end{tabular}
\end{table}

\subsection{误差分析/交叉验证}
<如适用>

\section{模型评价与推广}
\subsection{优点}
\begin{itemize}
    \item 具体做法 + 对本题解决了什么（动作化）。
\end{itemize}

\subsection{局限}
\begin{itemize}
    \item 具体局限 + 机制 + 后果 + 补救（机制化）。
\end{itemize}

\subsection{推广与落点}
面向决策者的可执行建议：对谁 + 做什么 + 预期效果 + 数据支撑。

\section{结论}
简要回顾各问核心结论与量化结果。

\bibliographystyle{gbt7714-2015}
\bibliography{references}

\appendix
\section{核心代码}
\lstinputlisting[language=Python]{../src/q1.py}
% \lstinputlisting[language=Python]{../src/q2.py}

\section{数据清洗与预处理说明}
<简述清洗步骤、缺失填补、异常值处理等>

\end{document}
"""

# DOCX 模板（美赛 MCM/ICM，英文，25 页硬上限）
DOCX_TEMPLATE_NOTE = """# DOCX 模板说明

由于 Python 标准库不支持直接生成复杂 DOCX，请按以下结构在 Word 中手动创建或使用 `python-docx` 在脚本中生成。

## 美赛论文结构（25 页硬上限）

### 1. Summary Sheet（单独一页，不计入 25 页）
- Problem: <一句话>
- Approach: <一句话总思路>
- Key Results: 每问 1–2 句，Model + 关键数字（带单位/置信）
- Sensitivity / Strengths / Limitations: 各一句

### 2. 正文结构（含目录、参考文献、附录共 ≤25 页）

| 章节 | 内容要点 | 页数建议 |
|------|----------|----------|
| Contents | 目录（Word 自动生成） | 1 |
| 1. Introduction | 背景统计 + 引用 [1] 开头；问题重述 | 1–1.5 |
| 2. Assumptions | 每条带小标题（如 2.1 As a disease） | 1 |
| 3. Model Development | 每问一节：Model Name + Equations + Algorithm + Results | 8–12 |
| 4. Model Validation | Sensitivity + Error + Comparison | 2–3 |
| 5. Discussion | Strengths (actionable) / Limitations (mechanistic) | 1.5 |
| 6. Conclusions & Policy | Policy Letter / 给利益相关方的信 | 1 |
| References | 规范著录 | 0.5 |
| Appendix | 代码 + 数据说明 | 剩余 |

### 3. 写作规范
- 英文术语规范、句式清晰
- 图表三件套：引出句 → 图表（≥300dpi） → 解读句（量化结论）
- 所有数字来自 `out/results.json`
- 引用键与文末表一一对应

### 4. 生成建议
可在 `src/export_results.py` 中补充生成 `report/results_for_paper.json`，
再写一个小脚本用 `python-docx` 填充模板。
"""

# Python-docx 生成脚本模板
DOCX_GEN_TEMPLATE = '''#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""用 python-docx 生成美赛论文 DOCX 模板（需 pip install python-docx）。

用法: python gen_docx_template.py <输出路径>
"""
import sys
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_chinese_font(run, size=Pt(10.5), bold=False):
    run.font.size = size
    run.font.bold = bold
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')


def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei' if level <= 2 else 'SimSun')
    return h


def main(output_path):
    doc = Document()

    # 页面设置
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # 正文默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.25
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    # Summary Sheet
    add_heading_styled(doc, 'Summary Sheet', 1)
    doc.add_paragraph('Problem: <一句话描述问题>')
    doc.add_paragraph('Approach: <一句话总思路>')
    doc.add_paragraph('Key Results:')
    doc.add_paragraph('  Q1: <Model Name> — <关键数字，带单位/置信>')
    doc.add_paragraph('  Q2: <Model Name> — <关键数字，带单位/置信>')
    doc.add_paragraph('  Q3: <Model Name> — <关键数字，带单位/置信>')
    doc.add_paragraph('Sensitivity: <一句话>')
    doc.add_paragraph('Strengths: <一句话>')
    doc.add_paragraph('Limitations: <一句话>')

    doc.add_page_break()

    # Contents（Word 自动生成，此处占位）
    add_heading_styled(doc, 'Contents', 1)
    doc.add_paragraph('[右键更新域 → 更新整个目录]')

    doc.add_page_break()

    # 1. Introduction
    add_heading_styled(doc, '1. Introduction', 1)
    doc.add_paragraph('<背景统计 + 引用 [1] 开头>')
    doc.add_paragraph('<问题重述，用自己的话>')

    # 2. Assumptions
    add_heading_styled(doc, '2. Assumptions', 1)
    add_heading_styled(doc, '2.1 As a disease', 2)
    doc.add_paragraph('<假设内容 + 理由>')
    add_heading_styled(doc, '2.2 The Markovian assumption', 2)
    doc.add_paragraph('<假设内容 + 理由>')

    # 3. Model Development
    add_heading_styled(doc, '3. Model Development', 1)
    for i in range(1, 4):
        add_heading_styled(doc, f'3.{i} Question {i}: <问题名>', 2)
        add_heading_styled(doc, 'Model Formulation', 3)
        doc.add_paragraph('<模型名 + 核心公式>')
        add_heading_styled(doc, 'Solution Method', 3)
        doc.add_paragraph('<算法描述>')
        add_heading_styled(doc, 'Results', 3)
        doc.add_paragraph('<关键数值 + 图表引用>')

    # 4. Model Validation
    add_heading_styled(doc, '4. Model Validation', 1)
    add_heading_styled(doc, '4.1 Sensitivity Analysis', 2)
    doc.add_paragraph('<参数摄动表 + 灵敏度图>')
    add_heading_styled(doc, '4.2 Error Analysis', 2)
    doc.add_paragraph('<误差指标 + 交叉验证结果>')
    add_heading_styled(doc, '4.3 Comparison with Benchmarks', 2)
    doc.add_paragraph('<基准对比表>')

    # 5. Discussion
    add_heading_styled(doc, '5. Discussion', 1)
    add_heading_styled(doc, '5.1 Strengths', 2)
    doc.add_paragraph('<优点：具体做法 + 解决了什么>')
    add_heading_styled(doc, '5.2 Limitations', 2)
    doc.add_paragraph('<局限：具体局限 + 机制 + 后果 + 补救>')

    # 6. Conclusions & Policy
    add_heading_styled(doc, '6. Conclusions and Policy Recommendations', 1)
    doc.add_paragraph('<核心结论回顾>')
    add_heading_styled(doc, 'Policy Letter', 2)
    doc.add_paragraph('<面向非技术决策者的信，措辞通俗、可执行>')

    # References
    add_heading_styled(doc, 'References', 1)
    doc.add_paragraph('[1] Author. Title. Journal, Year.')
    doc.add_paragraph('[2] ...')

    doc.add_page_break()

    # Appendix
    add_heading_styled(doc, 'Appendix A: Core Code', 1)
    doc.add_paragraph('<核心代码或引用附件>')
    add_heading_styled(doc, 'Appendix B: Data Description', 1)
    doc.add_paragraph('<数据清洗、预处理说明>')

    doc.save(output_path)
    print(f"[gen_docx] 已生成 DOCX 模板: {output_path}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python gen_docx_template.py <output_path>", file=sys.stderr)
        sys.exit(1)
    try:
        import docx
    except ImportError:
        print("[err] 需安装 python-docx: pip install python-docx", file=sys.stderr)
        sys.exit(1)
    main(sys.argv[1)
'''

def write_file(path, content):
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)


def main():
    ap = argparse.ArgumentParser(description="生成数模解题标准目录骨架 + 论文模板")
    ap.add_argument("name", help="题名缩写，如 cumcm2026A")
    ap.add_argument("--dir", default=".", help="父目录，默认当前目录")
    ap.add_argument("--template", choices=["cumcm", "mcm", "both"], default="both",
                    help="生成哪种论文模板：cumcm(LaTeX)、mcm(DOCX)、both(默认)")
    args = ap.parse_args()

    root = os.path.join(args.dir, args.name)
    if os.path.exists(root):
        print(f"[init] 目标已存在，中止: {root}", file=sys.stderr)
        return 1

    dirs = ["data", "src", "out", os.path.join("report", "figures")]
    for d in dirs:
        os.makedirs(os.path.join(root, d), exist_ok=True)

    # 核心文件
    write_file(os.path.join(root, "src", "export_results.py"), EXPORT_TEMPLATE)
    write_file(os.path.join(root, "requirements.txt"), REQUIREMENTS)

    # README
    from datetime import datetime
    write_file(os.path.join(root, "README.md"),
               README_TEMPLATE.format(name=args.name, ts=datetime.now().isoformat(timespec="seconds")))

    # 论文模板
    report_dir = os.path.join(root, "report")
    if args.template in ("cumcm", "both"):
        write_file(os.path.join(report_dir, "main.tex"), LATEX_TEMPLATE)
        print(f"[init] 已生成 LaTeX 模板: {os.path.join(report_dir, 'main.tex')}")
    if args.template in ("mcm", "both"):
        write_file(os.path.join(report_dir, "main.docx"), "")  # 占位，实际用 gen_docx_template.py 生成
        write_file(os.path.join(root, "src", "gen_docx_template.py"), DOCX_GEN_TEMPLATE)
        print(f"[init] 已生成 DOCX 生成脚本: {os.path.join(root, 'src', 'gen_docx_template.py')}")
        print(f"       运行: python src/gen_docx_template.py report/main.docx")

    # 占位文件，避免空目录在某些工具里被忽略
    for rel in ["data/.gitkeep", "out/.gitkeep", "report/figures/.gitkeep"]:
        write_file(os.path.join(root, rel), "")

    print(f"[init] 已生成项目骨架: {root}")
    print("        data/ src/ out/ report/figures/")
    print("        requirements.txt + README.md")
    if args.template in ("cumcm", "both"):
        print("        report/main.tex (LaTeX, 国赛)")
    if args.template in ("mcm", "both"):
        print("        report/main.docx + src/gen_docx_template.py (DOCX, 美赛)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())